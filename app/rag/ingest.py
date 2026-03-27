import ipaddress
import socket
import structlog
import httpx
from pathlib import Path
from urllib.parse import urlparse
from bs4 import BeautifulSoup

from app.config import get_settings
from app.database import get_session, Document, init_db
from app.llm.provider import get_llm_provider

logger = structlog.get_logger()

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".txt", ".csv"}


def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    """Split text into overlapping chunks."""
    words = text.split()
    chunks = []
    start = 0
    while start < len(words):
        end = start + chunk_size
        chunk = " ".join(words[start:end])
        if chunk.strip():
            chunks.append(chunk.strip())
        start += chunk_size - overlap
    return chunks


async def ingest_text(content: str, source: str, title: str = None) -> int:
    """Ingest plain text into the knowledge base."""
    settings = get_settings()
    llm = get_llm_provider()
    session = get_session()
    count = 0

    try:
        chunks = chunk_text(content, settings.chunk_size, settings.chunk_overlap)
        for i, chunk in enumerate(chunks):
            embedding = await llm.get_embedding(chunk)
            doc = Document(
                source=source,
                title=title or source,
                content=chunk,
                embedding=embedding,
                chunk_index=i,
            )
            session.add(doc)
            count += 1

        session.commit()
        logger.info("ingested_text", source=source, chunks=count)
        return count
    except Exception as e:
        session.rollback()
        logger.error("ingest_text_error", error=str(e))
        raise
    finally:
        session.close()


def _is_safe_url(url: str) -> bool:
    """#2: Validate URL to prevent SSRF — reject private/internal IPs."""
    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https"):
        return False
    if not parsed.hostname:
        return False
    try:
        resolved = socket.getaddrinfo(parsed.hostname, None)
        for entry in resolved:
            ip = ipaddress.ip_address(entry[4][0])
            if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_reserved:
                return False
    except (socket.gaierror, ValueError):
        return False
    return True


async def ingest_url(url: str, title: str = None) -> int:
    """Fetch and ingest content from a URL."""
    if not _is_safe_url(url):
        raise ValueError(f"URL not allowed (private/internal address): {url}")

    try:
        # #13: Use async httpx instead of sync requests
        async with httpx.AsyncClient() as client:
            response = await client.get(url, timeout=30, headers={
                "User-Agent": "WootBot/1.0 Knowledge Ingester"
            })
            response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")

        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()

        text = soup.get_text(separator="\n", strip=True)
        page_title = title or (soup.title.string if soup.title else url)

        return await ingest_text(text, source=url, title=page_title)
    except ValueError:
        raise
    except Exception as e:
        logger.error("ingest_url_error", url=url, error=str(e))
        raise


async def ingest_pdf(filepath: str, title: str = None) -> int:
    """Ingest a PDF file."""
    import pdfplumber
    text_parts = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

    full_text = "\n\n".join(text_parts)
    return await ingest_text(full_text, source=filepath, title=title or Path(filepath).name)


async def ingest_docx(filepath: str, title: str = None) -> int:
    """Ingest a DOCX file."""
    import docx
    doc = docx.Document(filepath)
    text_parts = [para.text for para in doc.paragraphs if para.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    full_text = "\n\n".join(text_parts)
    return await ingest_text(full_text, source=filepath, title=title or Path(filepath).name)


async def ingest_markdown(filepath: str, title: str = None) -> int:
    """Ingest a Markdown file."""
    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()

    if not title:
        for line in content.split("\n"):
            if line.startswith("# "):
                title = line.lstrip("# ").strip()
                break

    return await ingest_text(content, source=filepath, title=title or Path(filepath).name)


async def ingest_csv(filepath: str, title: str = None) -> int:
    """Ingest a CSV file — each row becomes searchable text."""
    import csv
    text_parts = []
    with open(filepath, "r", encoding="utf-8") as f:
        reader = csv.reader(f)
        headers = next(reader, None)
        if headers:
            for row in reader:
                row_text = " | ".join(f"{h}: {v}" for h, v in zip(headers, row) if v.strip())
                if row_text:
                    text_parts.append(row_text)

    full_text = "\n".join(text_parts)
    return await ingest_text(full_text, source=filepath, title=title or Path(filepath).name)


async def ingest_file(filepath: str, title: str = None) -> int:
    """Auto-detect file type and ingest."""
    ext = Path(filepath).suffix.lower()
    handlers = {
        ".pdf": ingest_pdf,
        ".docx": ingest_docx,
        ".md": ingest_markdown,
        ".txt": _ingest_txt,
        ".csv": ingest_csv,
    }
    handler = handlers.get(ext)
    if not handler:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {', '.join(SUPPORTED_EXTENSIONS)}")
    return await handler(filepath, title)


async def _ingest_txt(filepath: str, title: str = None) -> int:
    with open(filepath, "r", encoding="utf-8") as f:
        content = f.read()
    return await ingest_text(content, source=filepath, title=title or Path(filepath).name)


def get_document_stats() -> dict:
    """Get stats about ingested documents."""
    session = get_session()
    try:
        from sqlalchemy import func, distinct
        total_chunks = session.query(Document).count()
        total_sources = session.query(func.count(distinct(Document.source))).scalar()
        sources = session.query(
            Document.source, Document.title, func.count(Document.id).label("chunks")
        ).group_by(Document.source, Document.title).all()

        return {
            "total_chunks": total_chunks,
            "total_sources": total_sources,
            "sources": [
                {"source": s[0], "title": s[1], "chunks": s[2]}
                for s in sources
            ],
        }
    finally:
        session.close()


async def ingest_resolved_tickets(max_pages: int = 10) -> int:
    """Ingest resolved Chatwoot conversations as knowledge base entries.

    Fetches resolved tickets, formats each as a Q&A pair, and stores them
    so the bot can learn from successfully resolved conversations.
    """
    from app.chatwoot.client import ChatwootClient

    client = ChatwootClient()
    total_chunks = 0
    seen_ids = set()

    for page in range(1, max_pages + 1):
        data = await client.list_resolved_conversations(page=page)
        conversations = data.get("data", {}).get("payload", [])

        if not conversations:
            break

        for conv in conversations:
            conv_id = conv.get("id")
            if not conv_id or conv_id in seen_ids:
                continue
            seen_ids.add(conv_id)

            messages = await client.get_messages(conv_id)
            if len(messages) < 2:
                continue

            # Build a readable transcript
            lines = []
            for msg in messages:
                role = "Customer" if msg["role"] == "user" else "Agent"
                lines.append(f"{role}: {msg['content']}")

            transcript = "\n".join(lines)
            source = f"chatwoot-ticket-{conv_id}"
            title = f"Resolved ticket #{conv_id}"

            count = await ingest_text(transcript, source=source, title=title)
            total_chunks += count

        logger.info("ingested_tickets_page", page=page, conversations=len(conversations))

    logger.info("ingested_resolved_tickets", total_chunks=total_chunks, tickets=len(seen_ids))
    return total_chunks


def delete_document_by_source(source: str) -> int:
    """Delete all chunks from a specific source."""
    session = get_session()
    try:
        count = session.query(Document).filter(Document.source == source).delete()
        session.commit()
        logger.info("deleted_document", source=source, chunks_deleted=count)
        return count
    except Exception as e:
        session.rollback()
        logger.error("delete_error", error=str(e))
        raise
    finally:
        session.close()


# CLI entrypoint
if __name__ == "__main__":
    import asyncio
    import sys

    init_db()

    if len(sys.argv) < 2:
        print("Usage:")
        print("  python -m app.rag.ingest url https://docs.example.com 'Title'")
        print("  python -m app.rag.ingest file /path/to/doc.pdf 'Title'")
        print("  python -m app.rag.ingest tickets [max_pages]")
        print(f"  Supported files: {', '.join(SUPPORTED_EXTENSIONS)}")
        sys.exit(1)

    mode = sys.argv[1]

    async def main():
        if mode == "tickets":
            max_pages = int(sys.argv[2]) if len(sys.argv) > 2 else 10
            count = await ingest_resolved_tickets(max_pages=max_pages)
        elif len(sys.argv) < 3:
            print(f"Mode '{mode}' requires additional arguments.")
            sys.exit(1)
        elif mode == "url":
            title = sys.argv[3] if len(sys.argv) > 3 else None
            count = await ingest_url(sys.argv[2], title)
        elif mode == "file":
            title = sys.argv[3] if len(sys.argv) > 3 else None
            count = await ingest_file(sys.argv[2], title)
        elif mode == "text":
            title = sys.argv[3] if len(sys.argv) > 3 else None
            count = await ingest_text(sys.argv[2], source="manual", title=title)
        else:
            print(f"Unknown mode: {mode}. Use: url, file, text, tickets")
            sys.exit(1)
        print(f"Ingested {count} chunks")

    asyncio.run(main())
